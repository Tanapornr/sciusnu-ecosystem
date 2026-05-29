from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "digital-learning-ecosystem-report.md"
OUT_PATH = ROOT / "digital-learning-ecosystem-report.docx"

FONT_NAME = "TH SarabunPSK"


def set_run_font(run, size=16, bold=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_font(paragraph, size=16, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 24 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, 16, bold=bold)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    styles = doc.styles
    for style_name in ["Normal", "Body Text", "List Paragraph"]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(16)

    heading_sizes = {
        "Title": 20,
        "Subtitle": 16,
        "Heading 1": 18,
        "Heading 2": 17,
        "Heading 3": 16,
        "Heading 4": 16,
    }
    for style_name, size in heading_sizes.items():
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        if style_name.startswith("Heading"):
            style.font.color.rgb = RGBColor(15, 31, 53)

    normal = styles["Normal"].paragraph_format
    normal.line_spacing = 1.15
    normal.space_after = Pt(6)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("รายงานการออกแบบระบบนิเวศการเรียนรู้ดิจิทัล")
    set_run_font(r, 20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("การออกแบบระบบนิเวศการเรียนรู้ดิจิทัลสำหรับการพัฒนาสมรรถนะดิจิทัลและ AI\nของเจ้าหน้าที่คณะวิทยาศาสตร์")
    set_run_font(r, 18, bold=True)

    for _ in range(5):
        doc.add_paragraph()

    meta = [
        "จัดทำโดย",
        "ชื่อ - นามสกุล ............................................................",
        "รหัสนิสิต ............................................................",
        "สาขาวิชา ............................................................",
        "รายวิชา การจัดการระบบนิเวศการเรียนรู้ดิจิทัล",
        "อาจารย์ผู้สอน ............................................................",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, 16, bold=(line == "จัดทำโดย"))
    doc.add_page_break()


TOC_ITEMS = [
    "1. ข้อมูลรายวิชา (Course Information)",
    "2. การวิเคราะห์ผู้เรียน (Learner Analysis)",
    "3. เป้าหมายการเรียนรู้ (Learning Goals)",
    "4. แนวคิดระบบนิเวศการเรียนรู้ดิจิทัล",
    "5. การออกแบบ Learning Ecosystem",
    "6. การออกแบบวิธีการจัดการเรียนรู้ (Pedagogy Design)",
    "7. การบูรณาการเทคโนโลยีและ AI",
    "8. Learning Flow / Learning Journey",
    "9. การประเมินผล (Assessment Design)",
    "10. ต้นแบบนวัตกรรม (Prototype)",
    "11. การวิเคราะห์เชิงจริยธรรมและผลกระทบ",
    "12. ความเป็นไปได้และความยั่งยืน",
    "13. Reflection",
    "14. สรุปแนวคิดสำคัญ",
]


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("สารบัญ")
    set_run_font(r, 16, bold=True)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(13)
    table.columns[1].width = Cm(2)
    set_cell_text(table.rows[0].cells[0], "หัวข้อ", bold=True)
    set_cell_text(table.rows[0].cells[1], "หน้า", bold=True)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9EAF7")
    for item in TOC_ITEMS:
        cells = table.add_row().cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], "........")
    doc.add_paragraph("หมายเหตุ: สามารถอัปเดตเลขหน้าใน Microsoft Word ได้หลังตรวจหน้าจริงของเอกสาร")
    set_paragraph_font(doc.paragraphs[-1], 16)
    doc.add_page_break()


def add_flow_diagram(doc):
    p = doc.add_paragraph()
    r = p.add_run("แผนผัง Learning Flow / Learning Journey")
    set_run_font(r, 16, bold=True)

    headers = [
        "1. ก่อนเรียน\nLogin + Pre-test",
        "2. ระหว่างเรียน\nดูวิดีโอให้จบ",
        "3. หลังเรียน\nPost-test + บันทึกคะแนน",
        "4. นอกห้องเรียน\nประวัติ/รายการโปรด",
        "5. Lifelong Learning\nเรียนต่อและพัฒนาต่อเนื่อง",
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(3.0), Cm(3.0), Cm(3.2), Cm(3.0), Cm(3.2)]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        set_cell_text(cell, headers[i], bold=True, color=(255, 255, 255))
        set_cell_shading(cell, ["2563EB", "059669", "F59E0B", "7C3AED", "0F1F35"][i])

    p = doc.add_paragraph("ลำดับการทำงานหลัก: เข้าสู่ระบบ → Pre-test → Video → Post-test → Progress/Certificate → Lifelong Learning")
    set_paragraph_font(p, 16)


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return
    table = doc.add_table(rows=1, cols=len(parsed[0]))
    table.style = "Table Grid"
    for i, value in enumerate(parsed[0]):
        set_cell_text(table.rows[0].cells[i], value, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in parsed[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row[: len(cells)]):
            set_cell_text(cells[i], value)


def add_image(doc, alt, rel_path):
    image_path = (ROOT / rel_path).resolve()
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(alt)
    set_run_font(r, 16, bold=True)
    doc.add_picture(str(image_path), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def clean_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("<br/>", " ")
    return text.strip()


def build_docx():
    doc = Document()
    style_doc(doc)
    add_cover(doc)
    add_toc(doc)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    table_buffer = []
    in_mermaid = False
    mermaid_inserted = False

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for line in lines:
        raw = line.rstrip()
        stripped = raw.strip()
        if not stripped:
            flush_table()
            continue
        if stripped == "---":
            flush_table()
            continue
        if stripped.startswith("```mermaid"):
            flush_table()
            in_mermaid = True
            if not mermaid_inserted:
                add_flow_diagram(doc)
                mermaid_inserted = True
            continue
        if in_mermaid:
            if stripped.startswith("```"):
                in_mermaid = False
            continue
        if stripped.startswith("|"):
            table_buffer.append(stripped)
            continue

        flush_table()
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = clean_inline(heading_match.group(2))
            if text == "รายงานการออกแบบระบบนิเวศการเรียนรู้ดิจิทัล":
                continue
            p = doc.add_heading(text, level=level)
            set_paragraph_font(p, 18 if level == 1 else 16, bold=True)
            continue

        p = doc.add_paragraph()
        r = p.add_run(clean_inline(stripped))
        set_run_font(r, 16)

    flush_table()

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Learning Ecosystem for Faculty of Science Staff")
        set_run_font(run, 14)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_docx()
